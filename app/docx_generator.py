import re
import io

from docx import Document
from docx.shared import Pt


def _add_formatted_text(paragraph, text: str):
    """Add text with **bold** formatting support."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """Convert markdown text to a DOCX file in memory.

    Returns:
        BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    for line in markdown_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
